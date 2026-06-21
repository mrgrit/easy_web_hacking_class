#!/usr/bin/env python3
"""
주차별 실습 워크북(.docx) 생성기.
contents/training/lab_weekNN.yaml 을 읽어 각 주차의 단계 수에 맞춘 기록표가 든
실습워크북_WeekNN.docx 를 downloads/ 에 만든다. flag·정답은 넣지 않는다(빈 양식).

실행:
    pip install python-docx pyyaml
    python3 downloads/build_docx.py        # 레포 루트에서
"""
from pathlib import Path
import yaml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
LABS = ROOT / "contents" / "training"
OUT = ROOT / "downloads"

WEEK_TITLES = {
    "01": "AI와 AI 에이전트, 그리고 윤리",
    "02": "리눅스·Claude Code 첫걸음 + 내 첫 웹사이트 + GitHub",
    "03": "웹의 작동 원리 + 직접 해보는 웹 해킹 (DVWA)",
    "04": "AI 에이전트와 함께하는 모의해킹 (NeoBank)",
    "05": "mini-CTF (MediForum + CTFd)",
}


def blank_line(doc, label):
    p = doc.add_paragraph()
    run = p.add_run(label + "  ")
    run.bold = True
    p.add_run("________________________________________________")


def build(week):
    lab = yaml.safe_load((LABS / f"lab_week{week}.yaml").read_text(encoding="utf-8"))
    steps = lab.get("steps", [])

    doc = Document()
    # 제목
    h = doc.add_heading(f"AI 웹 해킹 특강 — 실습 워크북", level=0)
    sub = doc.add_heading(f"Week {week} · {WEEK_TITLES.get(week,'')}", level=1)

    # 표지 정보
    doc.add_paragraph()
    for lbl in ["이름", "학교 / 반", "날짜"]:
        blank_line(doc, lbl)
    doc.add_paragraph("☐ 나는 ‘해킹 윤리 서약서’에 서명했다.")

    # 오늘의 목표
    doc.add_heading("1. 오늘의 목표 (읽고 한 줄로 옮겨 적기)", level=2)
    blank_line(doc, "→")

    # 새로 배운 용어
    doc.add_heading("2. 새로 배운 용어 3개", level=2)
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "용어"
    t.rows[0].cells[1].text = "내가 이해한 뜻 (한 줄)"
    for i in range(1, 4):
        t.rows[i].cells[0].text = ""
        t.rows[i].cells[1].text = ""

    # 실습 기록표
    doc.add_heading("3. 실습 기록표", level=2)
    doc.add_paragraph("각 단계를 해보고 결과를 적으세요. (정확한 명령·정답은 교과서를 보고 직접 채웁니다)")
    tbl = doc.add_table(rows=len(steps) + 1, cols=5)
    tbl.style = "Table Grid"
    hdr = ["단계", "무엇을 했나", "결과", "화면에서 본 것 / 메모", "막힌 점"]
    for j, htext in enumerate(hdr):
        c = tbl.rows[0].cells[j]
        c.text = htext
        for r in c.paragraphs[0].runs:
            r.bold = True
    for i, st in enumerate(steps, start=1):
        row = tbl.rows[i].cells
        cat = st.get("category", "")
        row[0].text = f"{st.get('order', i)}\n({cat})"
        row[1].text = ""
        row[2].text = "성공 ☐\n실패 ☐"
        row[3].text = ""
        row[4].text = ""

    # 우와 포인트 / 회고 / 윤리
    doc.add_heading("4. 오늘의 ‘우와!’ 포인트", level=2)
    blank_line(doc, "→")
    doc.add_heading("5. 한 줄 회고", level=2)
    blank_line(doc, "→")
    doc.add_heading("6. 윤리 체크", level=2)
    doc.add_paragraph("☐ 나는 허락된 표적(실습 환경)에서만 공격/점검을 했다.")

    out = OUT / f"실습워크북_Week{week}.docx"
    doc.save(str(out))
    print(f"  생성: {out.name}  (단계 {len(steps)}개)")


def build_pledge():
    doc = Document()
    doc.add_heading("🔐 보안 윤리 서약서", level=0)
    doc.add_paragraph("AI 웹 해킹 특강").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "해킹은 자물쇠를 따는 기술과 같습니다. 잘 쓰면 사람을 지키고, 잘못 쓰면 사람을 다치게 합니다. "
        "그래서 우리는 도구를 배우기 전에 약속부터 합니다. 강력한 기술일수록 약속이 먼저입니다. "
        "아래 5가지를 읽고 동의한다면 이름과 날짜를 적고 서명하세요. 이 서약서가 다음 모든 주차의 입장권입니다.")
    doc.add_heading("나의 서약", level=2)
    pledges = [
        ("① 허락된 곳에서만 한다.",
         "나는 이 특강이 연습용으로 만든 표적(DVWA·NeoBank·MediForum·CTFd)에서만 배운 기술을 쓴다."),
        ("② 남의 시스템은 건드리지 않는다.",
         "나는 허락받지 않은 다른 사람·학교·회사의 시스템에 절대 무단으로 접근하지 않는다."),
        ("③ 알게 된 약점은 악용하지 않고 책임 있게 알린다.",
         "나는 우연히 약점을 발견해도 나쁜 데 쓰지 않고, 고칠 수 있도록 제대로 알린다."),
        ("④ 배운 기술을 좋은 방향으로 쓴다.",
         "나는 이 기술을 누군가를 지키고 더 안전한 세상을 만드는 데 쓴다(화이트해커처럼)."),
        ("⑤ 위반 시 책임은 나에게 있음을 안다.",
         "나는 이 약속을 어겼을 때의 법적·윤리적 책임이 나 자신에게 있다는 것을 분명히 안다."),
    ]
    for head, body in pledges:
        p = doc.add_paragraph()
        p.add_run(head).bold = True
        doc.add_paragraph(body)
    doc.add_paragraph(
        "📖 참고 — 정보통신망법 제48조: 허락(정당한 접근권한) 없이 남의 정보통신망에 들어가면 "
        "처벌받는다. 합법과 범죄를 가르는 건 기술이 아니라 딱 한 단어, ‘허락’이다.")
    doc.add_heading("서명란", level=2)
    for lbl in ["이름", "학교 · 반", "날짜", "서명"]:
        blank_line(doc, lbl)
    doc.add_paragraph("보호자 확인(선택): 보호자 이름 ______________   서명 ______________")
    doc.add_paragraph("위 내용을 모두 이해했으며, 나는 이 약속을 지킬 것을 서약합니다.")
    out = OUT / "보안서약서.docx"
    doc.save(str(out))
    print(f"  생성: {out.name}")


if __name__ == "__main__":
    print("배부 자료(.docx) 생성 중...")
    build_pledge()
    for wk in WEEK_TITLES:
        build(wk)
    print("완료.")
